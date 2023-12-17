import math
import time
import base64
import sys
import os
import numpy as np
from docx import Document  # Import the python-docx library

new_depth_limit = 100000
sys.setrecursionlimit(new_depth_limit)


def read_file(file_path):
    _, file_extension = os.path.splitext(file_path.lower())
    if file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_extension == '.docx':
        return read_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def write_file(file_path, content):
    _, file_extension = os.path.splitext(file_path.lower())
    if file_extension == '.txt':
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
    elif file_extension == '.docx':
        write_docx_file(file_path, content)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def read_docx_file(file_path):
    doc = Document(file_path)
    text_content = ""
    for paragraph in doc.paragraphs:
        text_content += paragraph.text + "\n"
    return text_content


def write_docx_file(file_path, content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)


def generate_blocks(source_block, num_iterations, target_block_number, block_size=1024):
    intermediate_blocks = [source_block]
    intermediate_blocks_len = len(intermediate_blocks)

    def generate_block_recursive(block, remaining_iterations, operation_count=0):
        nonlocal intermediate_blocks, intermediate_blocks_len

        if remaining_iterations == 0 or intermediate_blocks_len > target_block_number:
            return operation_count

        xor_result = np.bitwise_xor.reduce(block, axis=0)
        block = np.bitwise_xor(block, xor_result)
        operation_count += len(block)

        intermediate_blocks.append(block)
        intermediate_blocks_len += 1
        return generate_block_recursive(block, remaining_iterations - 1, operation_count)

    chunk_size = min(block_size, num_iterations)
    for i in range(0, num_iterations, chunk_size):
        total_operations = generate_block_recursive(
            intermediate_blocks[-1], chunk_size)
        if intermediate_blocks_len > target_block_number:
            break

    return intermediate_blocks, total_operations


def calculate_accuracy(original_block, decrypted_block):
    correct_bits = np.sum(original_block == decrypted_block)
    total_bits = len(original_block)
    accuracy_percentage = (correct_bits / total_bits) * 100
    return accuracy_percentage


def calculate_hamming_distance(original_block, decrypted_block):
    hamming_distance = np.count_nonzero(original_block != decrypted_block)
    return hamming_distance


def encrypt(source_block, block_number, num_iterations):
    intermediate_blocks, total_operations = generate_blocks(
        source_block, num_iterations, block_number)
    if block_number < len(intermediate_blocks):
        return intermediate_blocks[block_number], total_operations
    else:
        return [], 0


def decrypt(final_block, block_number, num_iterations, block_size=1024):
    if block_number >= num_iterations:
        print("Decryption failed. Block number out of range.")
        return [], 0

    decrypted_blocks = [final_block]
    intermediate_blocks_len = len(decrypted_blocks)

    def generate_block_recursive(block, remaining_iterations, operation_count=0):
        nonlocal decrypted_blocks, intermediate_blocks_len

        if remaining_iterations == 0:
            return operation_count

        xor_result = np.bitwise_xor.reduce(block, axis=0)
        block = np.bitwise_xor(block, xor_result)
        operation_count += len(block)

        decrypted_blocks.append(block)
        intermediate_blocks_len += 1
        return generate_block_recursive(block, remaining_iterations - 1, operation_count)

    chunk_size = min(block_size, num_iterations - block_number)
    for i in range(0, num_iterations - block_number, chunk_size):
        total_operations = generate_block_recursive(
            decrypted_blocks[-1], chunk_size)
        if intermediate_blocks_len > block_number:
            break

    return decrypted_blocks, total_operations


def string_to_binary(string):
    return np.array([int(bit) for byte in string.encode('utf-8') for bit in f"{byte:08b}"], dtype=np.uint8)


def binary_to_string(binary_values):
    if len(binary_values) % 8 != 0:
        raise ValueError("Binary values length must be a multiple of 8 for proper decoding.")

    bytes_data = bytes(
        int(''.join(map(str, binary_values[i:i + 8])), 2) for i in range(0, len(binary_values), 8))
    return bytes_data.decode('utf-8', errors='ignore')


def main():
    input_file = 'input2.docx'  # Change to the actual input file path
    encrypted_output_file = 'encrypted.docx'
    decrypted_output_file = 'decrypted.docx'
    
    print(f'Input File: {input_file}')

    input_file_size = os.path.getsize(input_file)
    input_file_size_kb = input_file_size/1024

    print(f"The size of the file is {input_file_size_kb:.2f} KB")

    # Determine file format based on extension
    _, input_file_extension = os.path.splitext(input_file.lower())
    is_docx = input_file_extension == '.docx'

    input_file_content = read_file(input_file)

    source_block = np.array(list(string_to_binary(input_file_content)))

    num_iterations = 2 ** math.ceil(math.log2(len(source_block)))

    block_number = int(input("Enter the block number for encryption: "))

    start_time = time.time()
    encrypted_block, encryption_operations = encrypt(
        source_block, block_number, num_iterations)
    end_time = time.time()
    encryption_time = end_time - start_time

    encrypted_base64 = base64.b64encode(bytes(encrypted_block)).decode('utf-8')

    write_file(encrypted_output_file, encrypted_base64)

    if block_number < len(encrypted_block):
        print("Block matched!")

        print(f'Encryption Time: {encryption_time:.4f} seconds')
        print(f'Number of XOR Operations (Encryption): {encryption_operations}')

    start_time = time.time()
    decrypted_blocks, decryption_operations = decrypt(
        encrypted_block, block_number, num_iterations)
    end_time = time.time()
    decryption_time = end_time - start_time

    accuracy_percentage = calculate_accuracy(source_block, decrypted_blocks[-1])

    print(f'Accuracy Percentage: {accuracy_percentage:.2f}%')

    decrypted_file_content = binary_to_string(bytes(decrypted_blocks[-1]))
    write_file(decrypted_output_file, decrypted_file_content)

    print(f'Decryption Time: {decryption_time:.4f} seconds')
    print(f'Number of XOR Operations (Decryption): {decryption_operations}')

    print()

    print("Final Results")
    print("------------------------")
    print(f'Input File: {input_file}')
    print(f'Encryption Time: {encryption_time:.4f} seconds')
    print(f'Decryption Time: {decryption_time:.4f} seconds')
    print(f'Number of XOR Operations (Encryption): {encryption_operations}')
    print(f'Number of XOR Operations (Decryption): {decryption_operations}')
    print(f'Accuracy Percentage: {accuracy_percentage:.2f}%')

    print("Encryption and decryption completed.\n")


if __name__ == "__main__":
    main()
